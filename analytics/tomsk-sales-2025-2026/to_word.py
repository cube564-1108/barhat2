"""
Конвертация отчёта в формат Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Создаём документ
doc = Document()

# Заголовок
title = doc.add_heading('Анализ продаж Томск 2025-2026', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Дата и период
doc.add_paragraph('Период: Январь - Июль 2025 vs 2026')
doc.add_paragraph('Дата отчёта: 05.08.2026')
doc.add_paragraph()

# Сводка
doc.add_heading('Сводка', 1)

# Данные для таблицы
data = [
    ['Выручка', '36,5 млн руб.', '18,5 млн руб.', '-49.3%'],
    ['Заказы', '12 308', '5 572', '-54.7%'],
    ['Средний чек', '2 962 руб.', '3 318 руб.', '+12.0%']
]

# Таблица сводки
table = doc.add_table(rows=len(data) + 1, cols=4)
table.style = 'Light Grid Accent 1'

headers = ['Показатель', '2025', '2026', 'Изменение']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for i, row_data in enumerate(data):
    row = table.rows[i + 1]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

doc.add_paragraph()
doc.add_paragraph('Ключевой вывод: Бизнес показал существенное снижение продаж — почти в 2 раза по выручке и заказам. При этом средний чек вырос на 12%, что может указывать на смещение фокуса на более дорогие позиции.')

# Динамика по каналам
doc.add_page_break()
doc.add_heading('Динамика по каналам', 1)

# Каналы с ростом
doc.add_heading('Каналы с ростом продаж', 2)

growth_table = doc.add_table(rows=3, cols=4)
growth_table.style = 'Light Grid Accent 1'

growth_headers = ['Канал', '2025', '2026', 'Изменение']
for i, header in enumerate(growth_headers):
    cell = growth_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

growth_data = [
    ['МАКС', '0', '975 546 руб.', 'Новый канал'],
    ['Telegram', '586 160', '797 973 руб.', '+36.1%']
]

for i, row_data in enumerate(growth_data):
    row = growth_table.rows[i + 1]
    for j, value in enumerate(row_data):
        row.cells[j].text = value

doc.add_paragraph('Вывод: Telegram показывает рост и становится значимым каналом. МАКС — новый канал с потенциально хорошим стартом.')

# Каналы со снижением
doc.add_heading('Каналы со снижением продаж', 2)

declining_data = [
    ['WhatsApp', '3,23 млн руб.', '190 тыс. руб.', '-94%', '936', '48'],
    ['Ozon', '58 тыс. руб.', '5 тыс. руб.', '-92%', '21', '2'],
    ['Персонал', '77 тыс. руб.', '21 тыс. руб.', '-73%', '59', '23'],
    ['Купер', '465 тыс. руб.', '132 тыс. руб.', '-72%', '158', '40'],
    ['Яндекс Маркет', '547 тыс. руб.', '202 тыс. руб.', '-63%', '206', '59'],
    ['Flawery', '374 тыс. руб.', '143 тыс. руб.', '-62%', '153', '56'],
    ['Instagram', '466 тыс. руб.', '181 тыс. руб.', '-61%', '100', '33'],
    ['Улица', '4,96 млн руб.', '1,99 млн руб.', '-60%', '2 310', '752'],
    ['Заказ с сайта', '13,77 млн руб.', '5,59 млн руб.', '-59%', '3 741', '1 446'],
    ['Flowwow', '4,71 млн руб.', '3,11 млн руб.', '-34%', '1 924', '1 196'],
    ['Яндекс Еда', '4,65 млн руб.', '3,19 млн руб.', '-32%', '1 760', '1 072'],
    ['ВКонтакте', '726 тыс. руб.', '501 тыс. руб.', '-31%', '220', '104'],
    ['Чат на сайте', '341 тыс. руб.', '255 тыс. руб.', '-25%', '104', '75']
]

declining_table = doc.add_table(rows=len(declining_data) + 1, cols=7)
declining_table.style = 'Light Grid Accent 1'

declining_headers = ['Канал', 'Выручка 2025', 'Выручка 2026', 'Изменение', 'Заказы 2025', 'Заказы 2026', 'Изменение заказов']
for i, header in enumerate(declining_headers):
    cell = declining_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

for i, row_data in enumerate(declining_data):
    row = declining_table.rows[i + 1]
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    row.cells[2].text = row_data[2]
    row.cells[3].text = row_data[3]
    row.cells[4].text = row_data[4]
    row.cells[5].text = row_data[5]
    # Вычисляем изменение заказов (убираем пробелы)
    val_2025 = int(row_data[4].replace(' ', ''))
    val_2026 = int(row_data[5].replace(' ', ''))
    change = ((val_2026 - val_2025) / val_2025 * 100) if val_2025 > 0 else 0
    row.cells[6].text = f'{change:.1f}%'

# Группировка каналов
doc.add_paragraph()
doc.add_paragraph('Группировка каналов по степени снижения:')

doc.add_paragraph('Критическое снижение (>70%):', style='List Bullet')
doc.add_paragraph('WhatsApp (-94%), Ozon (-92%), Персонал (-73%), Купер (-72%)', style='List Bullet 2')

doc.add_paragraph('Сильное снижение (50-70%):', style='List Bullet')
doc.add_paragraph('Яндекс Маркет, Flawery, Instagram, Улица (-60%), Заказ с сайта (-59%)', style='List Bullet 2')

doc.add_paragraph('Умеренное снижение (25-50%):', style='List Bullet')
doc.add_paragraph('Flowwow (-34%), Яндекс Еда (-32%), ВКонтакте (-31%), Чат на сайте (-25%)', style='List Bullet 2')

# Динамика по товарам
doc.add_page_break()
doc.add_heading('Динамика по товарам', 1)

# Читаем данные по товарам
import pandas as pd
import re

df_data = pd.read_excel('Заказы Томск 25-26.xlsx')
df_data['Дата и время'] = pd.to_datetime(df_data['Дата и время'], format='%d.%m.%Y %H:%M')
df_full_data = df_data[df_data['Дата и время'].dt.to_period('M') < '2026-08']

def parse_compose(compose_str):
    if pd.isna(compose_str):
        return []
    items = compose_str.split(';')
    products = []
    for item in items:
        item = item.strip()
        if not item:
            continue
        qty_match = re.search(r'(\d+)\s*шт', item)
        quantity = int(qty_match.group(1)) if qty_match else 1
        without_price = re.sub(r'\s*-\s*[\d,.]+\s*₽.*$', '', item)
        without_qty = re.sub(r',?\s*\d+\s*шт\.?\s*$', '', without_price)
        name = re.sub(r'\s*\d+$', '', without_qty).strip()
        name = re.sub(r'\s*\(?\d+[а-яА-Яa-zA-Z]?\)?\s*$', '', name).strip()
        if name:
            products.append({'name': name, 'quantity': quantity})
    return products

all_products_data = []
for idx, row in df_full_data.iterrows():
    products = parse_compose(row['Состав'])
    for prod in products:
        all_products_data.append({
            'date': row['Дата и время'],
            'product_name': prod['name'],
            'quantity': prod['quantity']
        })

products_df_data = pd.DataFrame(all_products_data)
products_2025_data = products_df_data[products_df_data['date'].dt.year == 2025]
products_2026_data = products_df_data[products_df_data['date'].dt.year == 2026]

prod_2025_data = products_2025_data.groupby('product_name').agg({'quantity': 'sum'}).reset_index()
prod_2026_data = products_2026_data.groupby('product_name').agg({'quantity': 'sum'}).reset_index()

prod_comparison_data = []
all_prod_names_data = set(prod_2025_data['product_name']) | set(prod_2026_data['product_name'])

for prod in all_prod_names_data:
    row_2025 = prod_2025_data[prod_2025_data['product_name'] == prod]
    row_2026 = prod_2026_data[prod_2026_data['product_name'] == prod]
    qty_2025 = row_2025['quantity'].values[0] if not row_2025.empty else 0
    qty_2026 = row_2026['quantity'].values[0] if not row_2026.empty else 0
    if qty_2025 >= 50:
        change = ((qty_2026 - qty_2025) / qty_2025 * 100) if qty_2025 > 0 else 0
        prod_comparison_data.append({
            'name': prod,
            'qty_2025': qty_2025,
            'qty_2026': qty_2026,
            'change': change
        })

prod_comparison_data.sort(key=lambda x: x['change'])

# Топ-25 товаров со снижением
declining_products = [p for p in prod_comparison_data if p['change'] < -15][:25]

doc.add_heading('Топ-25 товаров со снижением продаж (>15%)', 2)

if declining_products:
    prod_table = doc.add_table(rows=len(declining_products) + 1, cols=4)
    prod_table.style = 'Light Grid Accent 1'

    # Заголовки
    prod_headers = ['Товар', '2025 (шт)', '2026 (шт)', 'Изменение']
    for i, header in enumerate(prod_headers):
        cell = prod_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Данные
    for i, p in enumerate(declining_products):
        row = prod_table.rows[i + 1]
        row.cells[0].text = p['name']
        row.cells[1].text = str(p['qty_2025'])
        row.cells[2].text = str(p['qty_2026'])
        row.cells[3].text = f'{p["change"]:.1f}%'

    doc.add_paragraph()
    doc.add_paragraph('Примечание: Товары с изменением -100% либо сняты с производства, либо переименованы.')
else:
    doc.add_paragraph('Нет товаров со снижением более 15%')

# Выводы
doc.add_page_break()
doc.add_heading('Выводы', 1)

doc.add_heading('1. Критическая ситуация', 2)
doc.add_paragraph('Business потерял почти половину выручки (-49%) и более половины заказов (-55%) за год. Это требует немедленного внимания.')

doc.add_heading('2. Каналы-проблемы', 2)
doc.add_paragraph('WhatsApp практически умер (-94%) — либо канал перестал работать, либо перестали отслеживаться заказы')
doc.add_paragraph('Улица и Заказ с сайта — основные каналы, показали снижение ~60%')

doc.add_heading('3. Точки роста', 2)
doc.add_paragraph('Telegram показывает рост (+36%) — стоит инвестировать в развитие')
doc.add_paragraph('МАКС — новый канал с потенциальным ростом')

doc.add_heading('4. Ассортимент', 2)
doc.add_paragraph('Множество позиций исчезло из продаж. Требуется ревизия ассортимента.')

# Рекомендации
doc.add_heading('Рекомендации', 1)

doc.add_heading('Срочные действия', 2)
doc.add_paragraph('1. Разобраться с WhatsApp — почему -94%? Это техническая проблема или реальный спад?', style='List Number')
doc.add_paragraph('2. Анализ "Улица" и "Заказ с сайта" — основная выручка падает, нужны маркетинговые усилия', style='List Number')
doc.add_paragraph('3. Ревизия ассортимента — почему столько товаров исчезло? вернуть популярное или объяснить падение', style='List Number')

doc.add_heading('Точки роста', 2)
doc.add_paragraph('4. Инвестировать в Telegram — канал растёт, нужно масштабировать', style='List Number')
doc.add_paragraph('5. Развивать МАКС — новый канал показал хороший старт', style='List Number')

doc.add_heading('Аналитика', 2)
doc.add_paragraph('6. Настроить регулярный мониторинг — помесячные отчёты по каналам и товарам', style='List Number')
doc.add_paragraph('7. Отслеживать средний чек — он растёт (+12%), это хороший сигнал', style='List Number')

# Футер
doc.add_paragraph()
doc.add_paragraph()
paragraph = doc.add_paragraph('Отчёт сформирован автоматически на основе данных за янв-июль 2025-2026')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph.runs[0].font.size = Pt(10)
paragraph.runs[0].font.color.rgb = RGBColor(128, 128, 128)

# Сохранение
doc.save('report.docx')
print('Отчёт сохранён в report.docx')
